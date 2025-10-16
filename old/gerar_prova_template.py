
import json
import random
from pathlib import Path
from typing import Dict, List, Tuple
from datetime import datetime
import argparse

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

DIFF_PERMITIDAS = {"média", "alta"}

def load_questions(json_path: Path):
    data = json.loads(json_path.read_text(encoding="utf-8"))
    # filtra dificuldades
    return [q for q in data if str(q.get("dificuldade","")).lower() in DIFF_PERMITIDAS]

def choose_balanced(questions_by_file: Dict[Path, List[dict]], total: int) -> List[Tuple[Path, dict]]:
    files = list(questions_by_file.keys())
    # Embaralha ordem de arquivos para não viciar sempre os mesmos
    random.shuffle(files)

    # Quantidade base por arquivo (pode ser 0 se total < len(files))
    base = max(total // len(files), 0)
    take = {f: min(base, len(questions_by_file[f])) for f in files}
    allocated = sum(take.values())

    # Sobra a distribuir
    remainder = total - allocated

    # Prioriza arquivos com mais disponibilidade
    candidates = sorted(files, key=lambda f: len(questions_by_file[f]) - take[f], reverse=True)
    for f in candidates:
        if remainder <= 0:
            break
        # tentar adicionar 1 por vez (no máx. +1 por ciclo)
        if len(questions_by_file[f]) > take[f]:
            take[f] += 1
            remainder -= 1

    # Se ainda faltou (pouca disponibilidade global), completará com o que existir
    chosen = []
    for f in files:
        pool = list(questions_by_file[f])
        random.shuffle(pool)
        n = min(take[f], len(pool))
        chosen.extend([(f, q) for q in pool[:n]])

    # Caso ainda tenhamos menos que 'total' (pouquíssima oferta), apenas retorna o que deu
    return chosen[:total]

def shuffle_alternatives(q: dict):
    alts = list(q.get("alternativas", []))
    correct = q.get("correta", "")
    options = alts + [correct]
    random.shuffle(options)
    # mapear letras
    letters = ["A","B","C","D","E"]
    mapped = list(zip(letters, options))
    correct_letter = next(l for l, t in mapped if t == correct)
    return mapped, correct_letter

def set_section_columns(section, num_cols=2, space_twips=708):
    sectPr = section._sectPr
    cols = sectPr.xpath("./w:cols")
    if cols:
        cols = cols[0]
    else:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num_cols))
    cols.set(qn("w:space"), str(space_twips))

def insert_questions_block(doc: Document, questions: List[Tuple[Path, dict]]):
    # Cria nova seção em 2 colunas para as questões
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec = doc.sections[-1]
    # margens discretas
    from docx.shared import Cm
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    set_section_columns(sec, 2)

    # adiciona questões
    answer_key = []
    for idx, (_f, q) in enumerate(questions, start=1):
        # enunciado
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {q.get('enunciado','').strip()}")
        run.font.size = Pt(10)

        # figuras
        imgs = []
        if isinstance(q.get("figura"), str):
            imgs.append(q["figura"])
        if isinstance(q.get("imagens"), list):
            imgs.extend(q["imagens"])
        for fig in imgs:
            fig_path = Path(fig)
            if not fig_path.is_absolute():
                # tenta relativo ao diretório de execução
                fig_path = Path.cwd() / fig_path
            if fig_path.exists():
                try:
                    doc.add_picture(str(fig_path), width=Inches(2.8))
                except Exception:
                    pass

        # afirmacoes
        if isinstance(q.get("afirmacoes"), dict) and q["afirmacoes"]:
            for key in sorted(q["afirmacoes"].keys(), key=lambda x: str(x)):
                sp = doc.add_paragraph()
                sr = sp.add_run(f"{key}) {q['afirmacoes'][key]}")
                sr.font.size = Pt(10)

        # alternativas (embaralhadas)
        mapped, letter = shuffle_alternatives(q)
        for letra, texto in mapped:
            ap = doc.add_paragraph()
            ar = ap.add_run(f"{letra}) {texto}")
            ar.font.size = Pt(10)

        answer_key.append(letter)
    return answer_key

def replace_placeholder_with_marker(template_doc: Document, placeholder: str) -> bool:
    """
    Procura um parágrafo com o texto exato do placeholder e o remove,
    retornando True se encontrado (indicando que a próxima inserção já está
    no lugar correto). Caso não encontre, retorna False.
    """
    for p in template_doc.paragraphs:
        if p.text.strip() == placeholder:
            # limpa o parágrafo (remove runs)
            for i in range(len(p.runs)-1, -1, -1):
                p._element.remove(p.runs[i]._r)
            # deixa um parágrafo "vazio" como âncora; as questões virão depois em seção nova
            return True
    return False

def build_from_template(template_path: Path, json_paths: List[Path], out_path: Path,
                        total_questions: int = 10, placeholder="{{QUESTOES}}", seed: int = None):
    if seed is not None:
        random.seed(seed)

    # carrega template
    doc = Document(str(template_path))

    # carrega bancos
    questions_by_file = {}
    for jp in json_paths:
        qs = load_questions(jp)
        random.shuffle(qs)
        questions_by_file[jp] = qs

    # seleciona balanceado
    selected = choose_balanced(questions_by_file, total_questions)

    # tenta posicionar no placeholder (mantém a "sequência do template")
    found = replace_placeholder_with_marker(doc, placeholder)

    # insere bloco de questões (2 colunas)
    answer_key = insert_questions_block(doc, selected)

    # Seção do gabarito (1 coluna)
    doc.add_section(WD_SECTION.NEW_PAGE)
    sec2 = doc.sections[-1]
    # voltar para 1 coluna
    set_section_columns(sec2, 1)

    p = doc.add_paragraph()
    r = p.add_run("GABARITO")
    r.bold = True
    r.font.size = Pt(12)

    gp = doc.add_paragraph()
    gp_run = gp.add_run(" | ".join(f"Q{i+1}: {answer_key[i]}" for i in range(len(answer_key))))
    gp_run.font.size = Pt(11)

    doc.save(str(out_path))

def gerar_prova_template(template, jsons, output, num=10, seed=None, placeholder="{{QUESTOES}}"):
    # chama a função build_from_template interna
    build_from_template(
        Path(template),
        [Path(j) for j in jsons],
        Path(output),
        total_questions=num,
        placeholder=placeholder,
        seed=seed
    )

if __name__ == "__main__":
    gerar_prova_template(
        template="/mnt/data/PROVA1_2025_1.docx",
        jsons=[
            "questoes1.json",
            "questoes2.json",
            "questoes3.json",
            "questoes4.json",
            "questoes5.json",
            "questoes6.json",
            "questoes7.json"
        ],
        output="/mnt/data/prova_via_template.docx",
        num=10,
        seed=123
    )

if __name__ == "__main__":
    gerar_prova_template(
        template="/mnt/data/PROVA1_2025_1.docx",
        jsons=[
            "questoes1.json",
            "questoes2.json",
            "questoes3.json",
            "questoes4.json",
            "questoes5.json",
            "questoes6.json",
            "questoes7.json"
        ],
        output="/mnt/data/prova_via_template.docx",
        num=10,
        seed=123
    )