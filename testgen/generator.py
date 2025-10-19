# -*- coding: utf-8 -*-
from __future__ import annotations
from typing import List, Dict, Any, Optional, Tuple
from pathlib import Path
from docx import Document
from docx.shared import Inches

from core import load_quiz

# -------------------------------
# Util
# -------------------------------

def mm_to_inches(mm: float) -> float:
    return (mm or 0) / 25.4

IMG_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".svg", ".pdf")

def _parse_img_spec(s: str) -> Tuple[str, Optional[float], Optional[float]]:
    """
    Interpreta 'path;LxA' -> (path, L, A) em mm; se não houver tamanho, retorna (path, None, None).
    """
    if not isinstance(s, str):
        return s, None, None
    if ';' in s:
        path, size = s.split(';', 1)
        import re
        m = re.match(r'^\s*(\d+(?:\.\d+)?)x(\d+(?:\.\d+)?)\s*$', size.strip())
        if m:
            return path.strip(), float(m.group(1)), float(m.group(2))
        return path.strip(), None, None
    return s.strip(), None, None

def _is_image_path(s: str) -> bool:
    if not isinstance(s, str):
        return False
    p = s.split(';', 1)[0].strip()
    return any(p.lower().endswith(ext) for ext in IMG_EXTS)

def _get_correta_tuple(q: Dict[str, Any]) -> Tuple[Optional[int], Any]:
    """
    Extrai a correta como (idx, valor). Se vier apenas o valor, idx será None.
    """
    cor = q.get("correta")
    if isinstance(cor, (list, tuple)) and len(cor) >= 2:
        try:
            idx = int(cor[0])
        except Exception:
            idx = None
        return idx, cor[1]
    return None, cor


# -------------------------------
# Carregamento (via core)
# -------------------------------

def _load_questions(json_paths: List[str], *, seed: Optional[int]) -> List[Dict[str, Any]]:
    """
    Lê cada JSON com o core (que normaliza, resolve e prepara alternativas) e
    anota o diretório base em '_base_dir' para resolver caminhos de imagem.
    """
    out: List[Dict[str, Any]] = []
    for p in json_paths:
        ds = load_quiz(p, seed=seed)
        base_dir = str(Path(p).parent.resolve())
        for q in ds.get("questions", []):
            if isinstance(q, dict):
                # não troca referência nem estrutura; só anota o base_dir para render
                if "_base_dir" not in q:
                    q["_base_dir"] = base_dir
                out.append(q)
    return out


# -------------------------------
# Renderização para DOCX
# -------------------------------

def _compose_docx_block(q: Dict[str, Any], seq: int) -> List[Dict[str, Any]]:
    """
    Transforma a questão em uma sequência de 'runs' para docx.
    - Numeração sequencial (1..N).
    - Alternativas rotuladas a), b), c)...
    - Imagens entre enunciado e alternativas.
    - Não faz resolve/shuffle local: tudo vem pronto do core.
    """
    runs: List[Dict[str, Any]] = []
    alph = "abcdefghijklmnopqrstuvwxyz"
    base_dir = q.get("_base_dir") or ""
    # Se quiser um placeholder padrão, defina aqui; por ora só um marcador textual quando não achar a imagem.

    # 1) Enunciado
    enun = str(q.get("enunciado") or "").strip()
    runs.append({"type": "text", "text": f"{seq}) {enun}\n"})

    # 2) Imagens do enunciado (entre enunciado e alternativas)
    imgs = q.get("imagens") or []
    if isinstance(imgs, list):
        for img in imgs:
            spec_p, wmm, hmm = _parse_img_spec(img)
            p = Path(base_dir, spec_p) if base_dir else Path(spec_p)
            runs.append({"type": "text", "text": ""})  # garante quebra/ancoragem
            if p.exists():
                runs.append({"type": "image", "path": str(p), "width_mm": wmm, "height_mm": hmm})
            else:
                runs.append({"type": "text", "text": "[imagem]\n"})
        if imgs:
            runs.append({"type": "text", "text": "\n"})  # respiro antes das alternativas

    # 3) Afirmacoes (se existirem)
    afirm = q.get("afirmacoes") or {}
    if isinstance(afirm, dict) and afirm:
        order = ["I","II","III","IV","V","VI","VII","VIII","IX","X"]
        for k in order:
            if k in afirm:
                runs.append({"type": "text", "text": f"  {k}. {str(afirm[k]).strip()}\n"})
        # Subenunciado (entre afirmativas e alternativas)
        sub = (q.get("subenunciado") or "").strip()
        if sub:
            runs.append({"type": "text", "text": f"  {sub}\n\n"})

    # 4) Alternativas (já preparadas pelo core: mescladas/deduplicadas/embaralhadas)
    alts = q.get("alternativas") or []
    for i, alt in enumerate(alts):
        label = alph[i] + ")" if i < len(alph) else f"{i+1})"
        s = str(alt or "")
        if _is_image_path(s):
            spec_p, wmm, hmm = _parse_img_spec(s)
            p = Path(base_dir, spec_p) if base_dir else Path(spec_p)
            runs.append({"type": "text", "text": f"  {label} "})
            if p.exists():
                runs.append({"type": "image", "path": str(p), "width_mm": wmm, "height_mm": hmm})
            else:
                runs.append({"type": "text", "text": "[imagem]\n"})
            runs.append({"type": "text", "text": "\n"})
        else:
            runs.append({"type": "text", "text": f"  {label} {s}\n"})

    runs.append({"type": "text", "text": "\n"})
    return runs

def _render_blocks_for_docx(resolved: List[Dict[str, Any]]) -> List[List[Dict[str, Any]]]:
    return [_compose_docx_block(q, i + 1) for i, q in enumerate(resolved)]


# -------------------------------
# Pipeline principal
# -------------------------------

def json2docx(
    json_paths: List[str],
    template: str,
    out_docx: str,
    placeholder: str = "{{QUESTOES}}",
    title: str = "Prova",
    num: Optional[int] = None,
    seed: Optional[int] = None,
    shuffle: bool = True,
) -> int:
    """
    Gera DOCX a partir de 1+ JSONs:
    - O core resolve variáveis (<VAR>, <VAR OP NUM>), normaliza chaves 'nome;valor' e
      prepara alternativas (merge correta, dedup, shuffle determinístico por questão).
    - Aqui embaralhamos **apenas** a ordem das questões (se 'shuffle=True').
    - Seleciona 'num' primeiras após o shuffle (se informado).
    - Insere figuras declaradas na questão (caminhos relativos ao JSON).
    """

    # 1) Carregar tudo via core (sem tratar "tipo")
    resolved: List[Dict[str, Any]] = _load_questions(json_paths, seed=seed)

    # 2) Embaralhar ordem das questões (opcional)
    if shuffle and len(resolved) > 1:
        import random
        rng = random.Random(seed)
        rng.shuffle(resolved)

    # 3) Selecionar N primeiras (se num>0)
    if isinstance(num, int) and num > 0:
        resolved = resolved[:num]

    # 4) Renderizar no DOCX (substituindo placeholder ou anexando ao fim)
    doc = Document(template)

    def replace_placeholder(document: Document, placeholder_text: str, blocks: List[List[Dict[str, Any]]]) -> bool:
        for para in document.paragraphs:
            if placeholder_text in para.text:
                p = para._p
                parent = p.getparent()
                idx = parent.index(p)
                parent.remove(p)
                # Inserir logo após o parágrafo removido
                for block in blocks:
                    pr = document.add_paragraph()
                    for run in block:
                        if run["type"] == "text":
                            pr.add_run(run["text"])
                        elif run["type"] == "image":
                            try:
                                wmm = run.get("width_mm"); hmm = run.get("height_mm")
                                if wmm and hmm:
                                    pr.add_run().add_picture(
                                        run["path"],
                                        width=Inches(mm_to_inches(wmm)),
                                        height=Inches(mm_to_inches(hmm)),
                                    )
                                else:
                                    pr.add_run().add_picture(run["path"], width=Inches(5.5))
                            except Exception:
                                pr.add_run("[imagem]")
                return True
        return False

    blocks = _render_blocks_for_docx(resolved)
    if not replace_placeholder(doc, placeholder, blocks):
        # Se placeholder não encontrado, adiciona ao final do documento
        for block in blocks:
            pr = doc.add_paragraph()
            for run in block:
                if run["type"] == "text":
                    pr.add_run(run["text"])
                elif run["type"] == "image":
                    try:
                        wmm = run.get("width_mm"); hmm = run.get("height_mm")
                        if wmm and hmm:
                            pr.add_run().add_picture(
                                run["path"],
                                width=Inches(mm_to_inches(wmm)),
                                height=Inches(mm_to_inches(hmm)),
                            )
                        else:
                            pr.add_run().add_picture(run["path"], width=Inches(5.5))
                    except Exception:
                        pr.add_run("[imagem]")

    Path(out_docx).parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_docx)
    return 0


# Backward compat para o seu GUI
def jsons_to_docx(
    json_paths,
    template,
    out_docx,
    placeholder='{{QUESTOES}}',
    title='Prova',
    num=None,
    seed=None,
    shuffle=True
):
    return json2docx(
        json_paths,
        template,
        out_docx,
        placeholder=placeholder,
        title=title,
        num=num,
        seed=seed,
        shuffle=shuffle
    )